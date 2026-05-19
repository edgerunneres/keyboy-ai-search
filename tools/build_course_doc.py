from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "KeyBoy搜索引擎课程设计增强版说明.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header_fill: str = "0F766E") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.margin_top = Cm(0.08)
            cell.margin_bottom = Cm(0.08)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(15, 118, 110 if level == 1 else 90)
    return paragraph


def add_body(doc: Document, text: str, bold_lead: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        rest = paragraph.add_run(text[len(bold_lead) :])
        runs = [lead, rest]
    else:
        runs = [paragraph.add_run(text)]
    for run in runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.15)
    section.right_margin = Cm(2.15)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    run = title.add_run("KeyBoy 3.0 LLM 多智能体在线研究系统说明")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Agentic RAG + 在线开放数据源 + LLM 证据合成 + Critic 校验")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(78, 91, 99)

    meta = doc.add_table(rows=5, cols=2)
    rows = [
        ("项目名称", "KeyBoy LLM 多智能体在线研究系统"),
        ("版本定位", "KeyBoy 3.0 / Agentic Deep Research 原型"),
        ("核心目标", "从本地搜索升级为可接入真实大模型的在线多智能体系统"),
        ("技术路线", "Agentic RAG、百炼/Qwen3.6、OpenAlex/arXiv 等在线源、BM25+语义排序、Critic 校验"),
        ("交付内容", "源代码、前端、在线源适配器、LLM Provider、系统设计、迭代说明、验收指南"),
    ]
    for idx, (key, value) in enumerate(rows):
        set_cell_text(meta.cell(idx, 0), key, bold=True)
        set_cell_text(meta.cell(idx, 1), value)
        set_cell_shading(meta.cell(idx, 0), "E6F4F1")
    meta.columns[0].width = Cm(4)
    meta.columns[1].width = Cm(11.5)

    doc.add_page_break()

    add_heading(doc, "1. 架构重构总览")
    add_body(
        doc,
        "KeyBoy 3.0 不再定位为本地小型搜索引擎，而是升级为 LLM 多智能体在线研究系统。系统采用 Agentic RAG / Deep Research 思路，将复杂问题拆解为规划、在线发现、证据索引、排序、合成和批判校验等步骤；有 API Key 时调用真实大模型，没有 API Key 时进入可解释 fallback。",
    )
    table = doc.add_table(rows=1, cols=3)
    headers = ["维度", "原计划", "增强版迭代"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
    diff_rows = [
        ("数据来源", "本地 JSON 小语料", "OpenAlex、Semantic Scholar、arXiv、Crossref 在线源 + 本地兜底"),
        ("智能体", "爬取、清洗、索引、搜索四类", "ResearchPlanner、OnlineDiscovery、EvidenceRanker、Synthesis、Critic"),
        ("大模型", "未接入真实模型", "OpenAI-compatible Chat Completions，可接入任意兼容模型"),
        ("输出", "搜索结果与摘要", "研究计划、在线证据、引用答案、风险校验、Agent Trace"),
        ("架构", "本地检索原型", "Agentic RAG / Deep Research 工作流"),
    ]
    for row in diff_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    style_table(table)

    add_heading(doc, "2. 核心创新点")
    add_bullets(
        doc,
        [
            "Agentic RAG：将研究问题拆解为规划、搜索、阅读、合成、批判多个阶段。",
            "在线开放数据源：接入 OpenAlex、Semantic Scholar、arXiv、Crossref。",
            "真实 LLM 接口：通过环境变量接入 OpenAI-compatible 大模型。",
            "证据引用：答案输出引用、来源、URL、时间和证据摘录。",
            "Critic 校验：检查是否使用真实模型、证据数量和来源多样性。",
            "稳定演示：没有 API Key 或网络受限时仍可用本地 fallback 运行。",
        ],
    )

    add_heading(doc, "3. 系统架构")
    arch = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["智能体", "职责", "验收价值"]):
        set_cell_text(arch.cell(0, idx), header, bold=True)
    arch_rows = [
        ("ResearchPlannerAgent", "理解问题、拆分子查询、选择在线源", "体现真实大模型 Agent 规划能力"),
        ("OnlineDiscoveryAgent", "访问 OpenAlex、Semantic Scholar、arXiv、Crossref", "不再局限本地小数据库"),
        ("CleanAgent", "清洗、去重、过滤低质量文本", "保证知识库质量"),
        ("IndexAgent", "构建 BM25 与语义向量索引", "体现搜索核心技术"),
        ("EvidenceRankerAgent", "排序在线证据，选择高相关资料", "提升答案证据质量"),
        ("SynthesisAgent", "调用 LLM 或 fallback 生成带引用答案", "输出研究报告式答案"),
        ("CriticAgent", "检查模型状态、证据数量与来源多样性", "提升可信度"),
    ]
    for row in arch_rows:
        cells = arch.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    style_table(arch, "2563EB")

    add_heading(doc, "4. 运行与验收")
    add_body(doc, "运行命令：python -m keyboy.app --host 127.0.0.1 --port 8787", "运行命令：")
    add_body(doc, "访问地址：http://127.0.0.1:8787", "访问地址：")
    add_body(doc, "真实模型配置：设置 DASHSCOPE_API_KEY 后默认启用百炼 OpenAI 兼容接口和 qwen3.6-max-preview。", "真实模型配置：")
    add_bullets(
        doc,
        [
            "查询“Agentic RAG GraphRAG LightRAG Self-RAG 最新研究怎么整合到课程项目”。",
            "展示 ResearchPlannerAgent 的子查询、OnlineDiscoveryAgent 的在线资料、SynthesisAgent 的答案。",
            "说明没有 API Key 时系统会明确显示 fallback，不伪装大模型调用。",
            "展示 CriticAgent 对证据不足、来源单一、未启用模型等风险的提示。",
        ],
    )

    add_heading(doc, "5. 当前验证结果")
    result = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["测试项", "结果", "说明", "验收意义"]):
        set_cell_text(result.cell(0, idx), header, bold=True)
    result_rows = [
        ("单元测试", "5/5 通过", "覆盖本地搜索与 Agentic Research 离线管线", "代码可复现"),
        ("在线源", "4 类", "OpenAlex、Semantic Scholar、arXiv、Crossref", "支撑在线大数据获取"),
        ("LLM 接口", "OpenAI-compatible", "通过环境变量配置模型服务", "可接入真实大模型"),
        ("Critic 校验", "已实现", "检查模型状态、证据数量、来源多样性", "提升可信度"),
    ]
    for row in result_rows:
        cells = result.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    style_table(result, "B7791F")

    add_heading(doc, "6. 推荐答辩表述")
    add_body(
        doc,
        "我们没有停留在本地 TF-IDF 搜索，而是把 KeyBoy 重构为 Agentic RAG / Deep Research 系统。现在它具备研究规划、在线开放数据源获取、证据排序、LLM 合成、引用输出和 Critic 风险校验。没有 API Key 时系统会明确进入 fallback，不会伪装调用大模型；有 API Key 时可接入任意 OpenAI-compatible 模型。这个架构已经从课程演示原型走向真实前沿系统的雏形。",
    )

    add_heading(doc, "7. 参考经验")
    add_bullets(
        doc,
        [
            "GraphRAG：从局部片段检索扩展到全局图谱 sensemaking。",
            "LightRAG：以轻量图结构提升 RAG 效率和上下文关联。",
            "Self-RAG：生成过程加入检索决策和自我批判。",
            "AutoGen / LangGraph：多智能体编排和可观测执行轨迹。",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
